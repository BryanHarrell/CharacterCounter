from docx import Document
import glob

totalLengthInCharacters = 0
lengthInCharacters = 0
documentProgress = Document()
documentProgress.add_paragraph()
paragraphsProgress = documentProgress.paragraphs

for files in (glob.glob("../Documents/*.docx")):
        document = Document(files)
        paragraphs = document.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                lengthInCharacters = lengthInCharacters + len(run.text.encode('utf8'))
                documentProgress.add_paragraph(run.text)

totalLengthInCharacters = lengthInCharacters + totalLengthInCharacters;
print ((totalLengthInCharacters)/5.0/1500)
documentProgress.save('demo.docx')
